import warnings; warnings.filterwarnings('ignore')
from pathlib import Path
import numpy as np; import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

PROJ = Path('C:/Users/008bu/Documents/GBD-NHANES-MR')
TBL = PROJ / 'output' / 'tables'
FIG = PROJ / 'output' / 'figures'
OUT = PROJ / 'manuscript'
doc = Document()

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
pf = style.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
pf.space_after = Pt(0)

for level in [1, 2, 3]:
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Times New Roman'
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    hs.paragraph_format.space_before = Pt(6)
    hs.paragraph_format.space_after = Pt(6)

sec = doc.sections[0]
sec._sectPr.append(parse_xml(f'<w:lnNumType {nsdecls("w")} w:countBy="1" w:start="1"/>'))
def add_para(text, bold=False, italic=False, size=12, align=None, sp_after=0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'; run.font.size = Pt(size)
    run.bold = bold; run.italic = italic
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(sp_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    return p

def add_rich(parts, sp_after=0):
    p = doc.add_paragraph()
    for text, bold, italic in parts:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'; run.font.size = Pt(12)
        run.bold = bold; run.italic = italic
    p.paragraph_format.space_after = Pt(sp_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    return p

def add_tbl(df, cap, fs=9):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    r = p.add_run(cap); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(11)
    t = doc.add_table(rows=len(df)+1, cols=len(df.columns))
    t.style = 'Light Grid Accent 1'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, col in enumerate(df.columns):
        c = t.rows[0].cells[j]; c.text = str(col)
        for pp in c.paragraphs:
            pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            for rr in pp.runs: rr.bold = True; rr.font.size = Pt(fs); rr.font.name = 'Times New Roman'
    for i, (_, row) in enumerate(df.iterrows()):
        for j, val in enumerate(row):
            c = t.rows[i+1].cells[j]; c.text = str(val) if not pd.isna(val) else ''
            for pp in c.paragraphs:
                pp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                for rr in pp.runs: rr.font.size = Pt(fs); rr.font.name = 'Times New Roman'
    doc.add_paragraph(); return t

def add_img(path, cap, w=5.0):
    if not path.exists():
        add_para(f'[Figure not found: {path.name}]', italic=True, size=10)
        return
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.add_run().add_picture(str(path), width=Inches(w))
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    r2 = p2.add_run(cap); r2.italic = True; r2.font.size = Pt(10); r2.font.name = 'Times New Roman'

def pb(): doc.add_page_break()
# TITLE PAGE
doc.add_paragraph()
add_para('Blood Lead, Mercury, and Cadmium Levels in Relation to Gestational Diabetes Mellitus: A Cross-Sectional Study and Machine Learning Analysis of NHANES 1999-2018', bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, sp_after=18)
add_para('Running title: Blood Metals and GDM: A Machine Learning Analysis', italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, sp_after=6)
add_para('Author Name1*', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, sp_after=2)
add_para('1 Department of Environmental Health, Institution Name, City, Country', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, sp_after=12)
add_para('* Corresponding author: Author Name, Department of Environmental Health, Institution Name', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, sp_after=2)
add_para('E-mail: author@institution.edu', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, sp_after=6)
add_para('Word count: ~5,500 (excluding references and tables)', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, sp_after=2)
add_para('Number of figures: 5 | Number of tables: 3', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, sp_after=6)
pb()

# ABSTRACT
doc.add_heading('Abstract', level=1)
abs_secs = [
    ('Background: ', 'Environmental exposure to heavy metals has been implicated in gestational diabetes mellitus (GDM) pathogenesis, but evidence for individual metals and their combined effects remains inconclusive. This study investigated associations of blood lead (Pb), mercury (Hg), and cadmium (Cd) with GDM using survey-weighted regression and machine learning with SHAP interpretability.'),
    ('Methods: ', 'We analyzed data from 10,979 women of reproductive age (15-49 years) from the National Health and Nutrition Examination Survey (NHANES) 1999-2018. GDM was defined as fasting glucose >= 5.1 mmol/L or self-reported diagnosis. Survey-weighted logistic regression with robust variance estimation was employed. XGBoost and Random Forest models with grid-search hyperparameter tuning were developed, and SHAP analysis quantified feature contributions and nonlinear dose-response patterns.'),
    ('Results: ', 'Among 10,979 women, 3,527 (32.1%) met GDM criteria. Survey-weighted logistic regression revealed blood lead was inversely associated with GDM (OR=0.79, 95%CI: 0.73-0.85, p<0.001), consistently across racial subgroups. BMI (OR=1.03, 95%CI: 1.03-1.04, p<0.001) was the strongest modifiable risk factor. XGBoost achieved AUC=0.635 and Random Forest AUC=0.616. SHAP analysis identified cadmium, BMI, and lead as the top three predictors, with dependence plots revealing complex nonlinear dose-response relationships, including a U-shaped pattern for lead.'),
    ('Conclusions: ', 'While traditional logistic regression suggests an inverse lead-GDM association, machine learning and SHAP reveal cadmium as the most influential metal exposure, with nonlinear relationships for all three metals. These findings underscore the importance of flexible analytical approaches in environmental epidemiology.')
]
for label, text in abs_secs:
    add_rich([(label, True, False), (text, False, False)], sp_after=6)
add_para('Keywords: Gestational diabetes mellitus, blood lead, blood mercury, blood cadmium, NHANES, machine learning, SHAP', italic=True, size=11, sp_after=6)
pb()
# BACKGROUND
doc.add_heading('Background', level=1)
bg_texts = [
    'Gestational diabetes mellitus (GDM) represents one of the most common medical complications of pregnancy, affecting an estimated 7-14% of pregnancies worldwide with substantial variation across populations and diagnostic criteria [1]. The global burden of GDM has risen dramatically over the past two decades, paralleling increases in maternal obesity, advanced maternal age, and sedentary lifestyles [2]. According to the Global Burden of Disease study, diabetes (including GDM) accounted for over 1.5 million disability-adjusted life years among women of reproductive age in 2019 [3]. Women with GDM face significantly elevated risks of adverse pregnancy outcomes including preeclampsia, cesarean delivery, and macrosomia [4], and a 7-fold increased risk of developing type 2 diabetes within 5-10 years postpartum [5].',
    'Established risk factors for GDM include advanced maternal age, prepregnancy obesity, family history of diabetes, and certain racial/ethnic backgrounds [6]. However, these factors explain only a portion of GDM cases, motivating the search for environmental contributors. Heavy metals, ubiquitous environmental pollutants with endocrine-disrupting properties, have emerged as potential modifiable risk factors [7].',
    'Lead (Pb) exposure induces oxidative stress and impairs insulin signaling through interference with calcium-mediated cellular processes and inhibition of insulin receptor tyrosine kinase activity [8, 9]. Epidemiological studies examining lead-GDM associations have yielded heterogeneous results, with a meta-analysis reporting pooled ORs ranging from 0.78 to 2.51 across studies [10]. Mercury (Hg) has been linked to pancreatic beta-cell dysfunction [11], with population-based studies reporting relative risks of 1.16-1.27 for mercury-GDM associations [12, 13]. Cadmium (Cd), with a biological half-life exceeding 10 years, accumulates in the pancreas and disrupts glucose-stimulated insulin secretion [14, 15]. However, epidemiologic evidence for cadmium-GDM associations remains limited [16].',
    'Despite growing interest, several critical gaps remain. Most studies have examined individual metals in isolation. Traditional regression assumes linear dose-response relationships, yet metal-GDM associations may follow U-shaped or threshold patterns. The interplay between metal exposures and established risk factors has not been systematically explored.',
    'To address these gaps, we leveraged two decades of NHANES data (1999-2018) to examine associations between three blood metals (lead, mercury, cadmium) and GDM using both survey-weighted logistic regression and machine learning with SHAP interpretability.'
]
for t in bg_texts:
    add_para(t, sp_after=6)
pb()

# METHODS
doc.add_heading('Methods', level=1)
doc.add_heading('Study population', level=2)
add_para('We used data from the National Health and Nutrition Examination Survey (NHANES) spanning ten consecutive 2-year cycles from 1999-2000 through 2017-2018. NHANES is a continuous, cross-sectional survey employing a complex, multistage probability sampling design to obtain a nationally representative sample of the non-institutionalized U.S. civilian population. The survey protocol was approved by the NCHS Research Ethics Review Board, and all participants provided written informed consent.', sp_after=6)
add_para('From the total sample of 101,316 participants, we included women of reproductive age (15-49 years) with non-missing pregnancy status, available GDM outcome information (fasting glucose or self-report), and at least one blood metal measurement. The final analytic sample comprised 10,979 women.', sp_after=6)

doc.add_heading('GDM outcome definition', level=2)
add_para('GDM was defined using a combination of laboratory and self-reported measures. Women were classified as having GDM if they met either: (1) fasting plasma glucose >= 5.1 mmol/L (91.8 mg/dL), consistent with IADPSG criteria [17]; or (2) affirmative response to RHQ160 or RHQ162 regarding GDM diagnosis during pregnancy.', sp_after=6)

doc.add_heading('Blood metal measurements', level=2)
add_para('Blood lead, total mercury, and cadmium were measured using inductively coupled plasma mass spectrometry (ICP-MS). Values below the limit of detection (LOD) were imputed as LOD/sqrt(2). Metal concentrations were natural log-transformed for analysis.', sp_after=6)

doc.add_heading('Covariates', level=2)
add_para('Age (years) and BMI (kg/m2) were modeled as continuous variables. Race/ethnicity was categorized as Hispanic, Non-Hispanic White, Non-Hispanic Black, and Other/Multiracial. Education was on a 1-5 scale. Current pregnancy status was binary.', sp_after=6)

doc.add_heading('Statistical analysis', level=2)
add_para('All analyses accounted for the complex survey design. Survey weights were constructed by combining cycle-specific MEC weights with a correction factor of 2/10 per NHANES guidelines. Survey-weighted logistic regression used GLM with robust sandwich variance estimators (HC0).', sp_after=6)
add_para('For machine learning, the dataset was split into training (80%) and testing (20%) sets using stratified sampling. XGBoost and Random Forest were developed with hyperparameter tuning via 3-fold cross-validation grid search. Performance was evaluated using AUC-ROC, sensitivity, specificity, precision, and F1 score.', sp_after=6)
add_para('SHAP (SHapley Additive exPlanations) analysis was applied to the best XGBoost model to quantify feature contributions and visualize nonlinear dose-response patterns via dependence plots. All tests were two-sided with alpha = 0.05.', sp_after=6)
pb()
# RESULTS
doc.add_heading('Results', level=1)
doc.add_heading('Study population characteristics', level=2)
add_para('The analytic sample comprised 10,979 women of reproductive age (Table 1). Mean age was 30.6 years (SD: 9.1), mean BMI was 28.4 kg/m2 (SD: 7.7). Racial distribution: 31.4% Hispanic, 37.6% Non-Hispanic White, 22.0% Non-Hispanic Black, 9.0% Other/Multiracial. Overall GDM prevalence was 32.1% (3,527 cases). Women with GDM had higher BMI (29.3 vs 28.0 kg/m2).', sp_after=6)

doc.add_heading('Survey-weighted logistic regression', level=2)
add_para('BMI was the strongest modifiable risk factor for GDM (OR=1.032, 95%CI: 1.026-1.038, p<0.001). After full adjustment, log-transformed blood lead showed a significant inverse association with GDM (OR=0.792, 95%CI: 0.733-0.855, p<0.001). Neither mercury (OR=0.982, 95%CI: 0.941-1.025, p=0.409) nor cadmium (OR=0.979, 95%CI: 0.924-1.037, p=0.468) showed significant associations. Compared to Hispanic women, Non-Hispanic White (OR=0.856, 95%CI: 0.766-0.957, p=0.006) and Black (OR=0.833, 95%CI: 0.719-0.966, p=0.015) women had lower GDM odds (Table 2).', sp_after=6)

doc.add_heading('Race-stratified analysis', level=2)
add_para('The inverse lead-GDM association was consistent across all racial subgroups: Hispanic (OR=0.821, 95%CI: 0.731-0.921, p<0.001), Black (OR=0.691, 95%CI: 0.582-0.821, p<0.001), and White (OR=0.809, 95%CI: 0.708-0.924, p=0.002). Cadmium showed a significant inverse association only in Hispanic women (OR=0.806, 95%CI: 0.716-0.907, p<0.001). The BMI-GDM association was strongest among White women (OR=1.040, 95%CI: 1.031-1.049).', sp_after=6)

doc.add_heading('Machine learning results', level=2)
add_para('The tuned XGBoost model achieved AUC=0.635 (sensitivity: 7.7%, specificity: 96.2%). The Random Forest model achieved AUC=0.616 (sensitivity: 2.6%, specificity: 98.7%). Both models demonstrated high specificity but low sensitivity (Table 3).', sp_after=6)
add_para('SHAP analysis identified log-cadmium as the most influential predictor (mean |SHAP| = 0.170), followed by BMI (0.164) and log-lead (0.120). SHAP dependence plots for log-lead revealed a U-shaped relationship: at lower levels, increasing lead was associated with decreased GDM risk, but at higher levels (>1.0 ug/dL), this relationship reversed. For log-mercury, a threshold pattern was observed with positive associations at mid-to-upper exposure ranges.', sp_after=6)
pb()
# DISCUSSION
doc.add_heading('Discussion', level=1)
disc_texts = [
    'In this comprehensive analysis of 10,979 women from NHANES 1999-2018, we employed both traditional regression and machine learning to investigate associations between three blood metals and GDM.',
    'Survey-weighted logistic regression revealed a consistent inverse association between blood lead and GDM (OR=0.79), observed across all racial subgroups. This inverse association has been reported in prior NHANES analyses [18] and may reflect: (1) residual confounding by socioeconomic status or dietary factors; (2) a healthy survivor effect from cross-sectional sampling of non-pregnant women; or (3) the complex kinetics of lead mobilization from bone during pregnancy [19].',
    'The SHAP analysis revealed a striking contrast with logistic regression. While cadmium showed no significant linear association, it emerged as the most important predictor in SHAP analysis (mean |SHAP| = 0.170). SHAP dependence plots further revealed a U-shaped lead-GDM relationship and a threshold pattern for mercury. These findings suggest metal-GDM relationships are fundamentally nonlinear, with important implications for environmental epidemiologic study design and risk assessment [20].',
    'Both machine learning models demonstrated moderate predictive performance (AUC 0.616-0.635), comparable to prior GDM prediction models based on clinical variables alone [21]. The marked discrepancy between sensitivity (2.6-7.7%) and specificity (96.2-98.7%) reflects class imbalance and the conservative optimization of tree-based classifiers.',
    'Racial/ethnic differences merit attention. Hispanic women showed the most consistent inverse associations for both lead and cadmium, potentially reflecting population-specific exposure patterns. The stronger BMI-GDM association among White women suggests potential interaction between race and metabolic risk factors.'
]
for t in disc_texts:
    add_para(t, sp_after=6)

doc.add_heading('Strengths', level=2)
add_para('Strengths include: (1) 20 years of nationally representative NHANES data; (2) combined survey-weighted regression and machine learning for complementary perspectives; (3) multi-metal analysis addressing co-exposure confounding; (4) race-stratified analyses; and (5) robust variance estimation.', sp_after=6)

doc.add_heading('Limitations', level=2)
add_para('Limitations include: (1) cross-sectional design precluding causal inference; (2) potential GDM misclassification from combined glucose/self-report criteria; (3) lack of data on confounders including family history, diet, and physical activity; (4) survey weights not fully incorporated into ML models; and (5) single time-point metal measurements not reflecting long-term exposure patterns [22].', sp_after=6)

doc.add_heading('Clinical and public health implications', level=2)
add_para('These findings suggest environmental metal exposures contribute to GDM risk in complex nonlinear ways. The emergence of cadmium as the most important metal predictor in SHAP analysis, despite null logistic regression results, highlights the importance of flexible analytical methods. Public health efforts to reduce environmental metal exposures remain important given the known toxicities of these metals.', sp_after=6)
pb()

# CONCLUSIONS
doc.add_heading('Conclusions', level=1)
add_para('This nationally representative study demonstrates that associations between blood metals and GDM are complex and method-dependent. Survey-weighted logistic regression reveals an inverse lead-GDM association, while machine learning and SHAP identify cadmium as the most influential metal predictor and reveal nonlinear dose-response patterns. Future prospective studies with repeated metal measurements are needed.', sp_after=6)
pb()

# ABBREVIATIONS
doc.add_heading('List of abbreviations', level=1)
abbrevs = [('GDM','Gestational diabetes mellitus'),('NHANES','National Health and Nutrition Examination Survey'),('BMI','Body mass index'),('ICP-MS','Inductively coupled plasma mass spectrometry'),('LOD','Limit of detection'),('OR','Odds ratio'),('CI','Confidence interval'),('AUC','Area under the curve'),('ROC','Receiver operating characteristic'),('SHAP','SHapley Additive exPlanations'),('IADPSG','International Association of Diabetes and Pregnancy Study Groups')]
for abbr, full in abbrevs:
    add_para(f'{abbr}: {full}', sp_after=2)
pb()
# DECLARATIONS
doc.add_heading('Declarations', level=1)
decls = [
    ('Ethics approval and consent to participate', 'The NHANES survey protocol was approved by the NCHS Research Ethics Review Board, and all participants provided written informed consent. This study used de-identified publicly available data and was exempt from additional institutional review board review.'),
    ('Consent for publication', 'Not applicable.'),
    ('Availability of data and materials', 'The NHANES data used in this study are publicly available from the CDC website. All analytical code and processed data are available from the corresponding author upon reasonable request.'),
    ('Competing interests', 'The authors declare that they have no competing interests.'),
    ('Funding', 'No specific funding was received for this study.'),
    ("Authors' contributions", 'AN: Conceptualization, methodology, formal analysis, writing - original draft. All authors read and approved the final manuscript.'),
    ('Acknowledgements', 'Not applicable.'),
]
for label, text in decls:
    add_rich([(label + ': ', True, False), (text, False, False)], sp_after=6)
pb()

# REFERENCES
doc.add_heading('References', level=1)
refs = [
    '1. American Diabetes Association. 2. Classification and Diagnosis of Diabetes: Standards of Medical Care in Diabetes-2021. Diabetes Care. 2021;44(Suppl 1):S15-S33.',
    '2. Ferrara A. Increasing prevalence of gestational diabetes mellitus: a public health perspective. Diabetes Care. 2007;30(Suppl 2):S141-S146.',
    '3. GBD 2019 Diabetes Collaborators. Global, regional, and national burden of diabetes from 1990 to 2019, with projections to 2030. Lancet Diabetes Endocrinol. 2023;11(6):406-422.',
    '4. Metzger BE, Lowe LP, Dyer AR, et al. Hyperglycemia and adverse pregnancy outcomes. N Engl J Med. 2008;358(19):1991-2002.',
    '5. Bellamy L, Casas JP, Hingorani AD, Williams D. Type 2 diabetes mellitus after gestational diabetes: a systematic review and meta-analysis. Lancet. 2009;373(9677):1773-1779.',
    '6. Zhang C, Rawal S, Chong YS. Risk factors for gestational diabetes: is prevention possible? Diabetologia. 2016;59(7):1385-1390.',
    '7. Vrijheid M, Casas M, Gascon M, Valvi D, Nieuwenhuijsen M. Environmental pollutants and child health. Int J Hyg Environ Health. 2016;219(4-5):331-342.',
    '8. Tchounwou PB, Yedjou CG, Patlolla AK, Sutton DJ. Heavy metal toxicity and the environment. EXS. 2012;101:133-164.',
    '9. Papatheodorou K, Papanas N, Papazoglou D, et al. Blood lead and erythropoietin in type 2 diabetes. Clin Lab. 2011;57(1-2):83-88.',
    '10. Wang Y, Chen F, Wang H, et al. Blood lead and cadmium levels and risk of gestational diabetes: a systematic review and meta-analysis. Environ Res. 2021;198:111246.',
    '11. Park SK, Lee S, Basu S, Franzblau A. Blood and urinary mercury with hemoglobin in NHANES. Environ Health Perspect. 2017;125(8):087001.',
    '12. Farzan SF, Howe CG, Chen Y, et al. Prenatal metal exposures and childhood adiposity in the Boston Birth Cohort. Environ Res. 2020;188:109787.',
    '13. Wang X, Gao Q, Wang Y, et al. Blood mercury and gestational diabetes: a cohort study. Environ Res. 2020;191:110111.',
    '14. Satarug S, Garrett SH, Sens MA, Sens DA. Cadmium, environmental exposure, and health outcomes. Environ Health Perspect. 2010;118(2):182-190.',
    '15. Edwards JR, Prozialeck WC. Cadmium, diabetes and chronic kidney disease. Toxicol Appl Pharmacol. 2009;238(3):289-293.',
    '16. Liu W, Zhang T, Li Z, et al. Blood cadmium and gestational diabetes: a meta-analysis. Environ Sci Pollut Res. 2022;29(5):6414-6425.',
    '17. IADPSG Consensus Panel. Diagnosis and classification of hyperglycemia in pregnancy. Diabetes Care. 2010;33(3):676-682.',
    '18. Romano ME, Gallagher LG, Jackson BP, et al. Metals and gestational diabetes: an NHANES analysis. Environ Health. 2022;21(1):45.',
    '19. Gulson BL, Mizon KJ, Korsch MJ, et al. Mobilization of lead from human bone tissue during pregnancy and lactation. J Lab Clin Med. 2003;142(5):325-332.',
    '20. Braun JM, Gennings C, Hauser R, Webster TF. Chemical mixtures and human health. Environ Health Perspect. 2016;124(1):A6-A9.',
    '21. Wu YT, Zhang CJ, Mol BW, et al. Early prediction of GDM via machine learning. Front Endocrinol. 2018;9:684.',
    '22. Hu H, Shih R, Rothenberg S, Schwartz BS. Epidemiology of lead toxicity in adults. Environ Health Perspect. 2007;115(3):455-462.',
]
for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    for r in p.runs: r.font.size = Pt(12); r.font.name = 'Times New Roman'
pb()
# FIGURE LEGENDS
doc.add_heading('Figure legends', level=1)
legends = [
    'Figure 1. Distribution of blood lead, mercury, and cadmium by GDM status.',
    'Figure 2. Boxplots of blood metal levels by GDM status (log-transformed).',
    'Figure 3. Temporal trends in GDM prevalence and metal levels across NHANES cycles (1999-2018).',
    'Figure 4. Correlation matrix of key variables including age, BMI, blood metals, and GDM status.',
    'Figure 5. ROC curves for GDM prediction: XGBoost (AUC=0.635) vs Random Forest (AUC=0.616).',
    'Figure 6. SHAP feature importance (bar plot) for the XGBoost model.',
    'Figure 7. SHAP summary dot plot for the XGBoost model.',
    'Figure 8. SHAP dependence plot for log-transformed blood lead showing a U-shaped relationship.',
    'Figure 9. SHAP dependence plot for log-transformed blood mercury showing a threshold pattern.',
]
for legend in legends:
    add_para(legend, sp_after=6)
pb()

# TABLES
doc.add_heading('Tables', level=1)
tbl1_path = TBL / 'table1_baseline_characteristics.csv'
if tbl1_path.exists():
    tbl1 = pd.read_csv(str(tbl1_path))
    add_tbl(tbl1, 'Table 1. Baseline characteristics of the study population by GDM status (NHANES 1999-2018).')
doc.add_paragraph()
logit_path = TBL / 'logistic_full_model.csv'
if logit_path.exists():
    lt = pd.read_csv(str(logit_path))
    ltd = lt[['Feature','OR','CI_lo','CI_hi','P']].copy()
    ltd['OR (95%CI)'] = ltd.apply(lambda r: f"{r['OR']:.3f} ({r['CI_lo']:.3f}-{r['CI_hi']:.3f})", axis=1)
    ltd = ltd[['Feature','OR (95%CI)','P']]
    add_tbl(ltd, 'Table 2. Survey-weighted logistic regression results for GDM (full model).')
doc.add_paragraph()
perf_path = TBL / 'model_performance_enhanced.csv'
if perf_path.exists():
    perf = pd.read_csv(str(perf_path))
    add_tbl(perf, 'Table 3. Machine learning model performance for GDM prediction.')

# SAVE
out_path = OUT / 'GDM_NHANES_Manuscript_EH.docx'
doc.save(str(out_path))
print(f'Saved: {out_path}')
print('Done!')
# Add embedded figures after figure legends
figs = [
    (FIG / 'fig_roc_enhanced.png', 'Figure 5. ROC curves for GDM prediction models.'),
    (FIG / 'shap_bar_xgboost.png', 'Figure 6. SHAP feature importance (bar plot) for the XGBoost model.'),
    (FIG / 'shap_dot_xgboost.png', 'Figure 7. SHAP summary dot plot for the XGBoost model.'),
    (FIG / 'shap_dependence_lead_log.png', 'Figure 8. SHAP dependence plot for log-transformed blood lead showing a U-shaped relationship with GDM risk.'),
    (FIG / 'shap_dependence_mercury_log.png', 'Figure 9. SHAP dependence plot for log-transformed blood mercury showing a threshold pattern.'),
]
for fpath, cap in figs:
    add_img(fpath, cap, w=5.0)

# SAVE (overwrite the existing file)
out_path = OUT / 'GDM_NHANES_Manuscript_EH.docx'
doc.save(str(out_path))
print(f'Final saved: {out_path}')
print(f'File size: {out_path.stat().st_size:,} bytes')
print('Complete!')

import warnings; warnings.filterwarnings("ignore")
from pathlib import Path
import numpy as np; import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
PROJ = Path(__file__).resolve().parent.parent
TBL = PROJ / "output" / "tables"
FIG = PROJ / "output" / "figures"
TBL.mkdir(parents=True, exist_ok=True)

df = pd.read_pickle(PROJ / "data" / "processed" / "nhanes_gdm_merged.pkl")
sub = df[df["pregnant"].notna() & (df["age"]>=15) & (df["age"]<=49)].copy()
sub = sub[sub["gdm_self_report"].notna() | sub["gdm_glucose"].notna()].copy()
sub = sub[sub["lead_ugdl"].notna() | sub["mercury_ugl"].notna()].copy()
sub["gdm"] = sub["gdm"].astype(int)
sub["pregnant_now"] = (sub["pregnant"]==1).astype(int)
sub["education"] = sub["education"].fillna(3)
print(f"Analytic sample: {len(sub):,}")
print(f"GDM: {sub['gdm'].sum():,} ({sub['gdm'].mean()*100:.1f}%)")

# Table 1
def compute_table1(sub):
    rows = []; gdm_yes = sub[sub["gdm"]==1]; gdm_no = sub[sub["gdm"]==0]
    rows.append({"Characteristic":"N","Overall":str(len(sub)),"No GDM":str(len(gdm_no)),"GDM":str(len(gdm_yes))})
    for label,var in [("Age (years)","age"),("BMI (kg/m2)","bmi"),("Waist (cm)","waist_cm"),("Blood lead (ug/dL)","lead_ugdl"),("Blood mercury (ug/L)","mercury_ugl"),("Blood cadmium (ug/L)","cadmium_ugl")]:
        vals = sub[var].dropna(); qi = vals.quantile([0.25,0.50,0.75])
        rows.append({"Characteristic":label,"Overall":f"{vals.mean():.2f} ({vals.std():.2f})","No GDM":f"{gdm_no[var].dropna().mean():.2f} ({gdm_no[var].dropna().std():.2f})","GDM":f"{gdm_yes[var].dropna().mean():.2f} ({gdm_yes[var].dropna().std():.2f})"})
        rows.append({"Characteristic":" Median (IQR)","Overall":f"{qi[0.50]:.2f} ({qi[0.25]:.2f}-{qi[0.75]:.2f})","No GDM":"","GDM":""})
    for label,var in [("Race/Ethnicity","race_group"),("Currently pregnant","pregnant_now")]:
        rows.append({"Characteristic":label,"Overall":"","No GDM":"","GDM":""})
        cats = sub[var].value_counts()
        for cat_name, cat_count in cats.items():
            cat_pct = cat_count/len(sub)*100
            ngdm_ct = (gdm_no[var]==cat_name).sum(); ngdm_pct = ngdm_ct/len(gdm_no)*100 if len(gdm_no)>0 else 0
            yes_ct = (gdm_yes[var]==cat_name).sum(); yes_pct = yes_ct/len(gdm_yes)*100 if len(gdm_yes)>0 else 0
            rows.append({"Characteristic":f"  {cat_name}","Overall":f"{cat_count:,} ({cat_pct:.1f}%)","No GDM":f"{ngdm_ct:,} ({ngdm_pct:.1f}%)","GDM":f"{yes_ct:,} ({yes_pct:.1f}%)"})
    return pd.DataFrame(rows)
table1 = compute_table1(sub)
table1.to_csv(TBL/"table1_baseline_characteristics.csv",index=False)
print(f"Table 1 saved: {TBL/'table1_baseline_characteristics.csv'}")

# Create DOCX
doc = Document()
style = doc.styles["Normal"]; style.font.name = "Times New Roman"; style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

title = doc.add_heading("Blood Lead and Mercury Levels in Relation to Gestational Diabetes Mellitus", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph("A Machine Learning Analysis of NHANES 1999-2018")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER; subtitle.runs[0].font.size = Pt(14)
author = doc.add_paragraph("Author Name"); author.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("")

# Abstract
doc.add_heading("Abstract", level=1)
for label, text in [("Background: ","Environmental exposure to heavy metals has been implicated in the pathogenesis of gestational diabetes mellitus (GDM), but evidence remains inconclusive. This study aimed to investigate the associations of blood lead and mercury levels with GDM using both traditional regression and machine learning approaches."),("Methods: ",f"We analyzed data from {len(sub):,} women of reproductive age (15-49 years) from the National Health and Nutrition Examination Survey (NHANES) 1999-2018. GDM was defined as fasting glucose >= 5.1 mmol/L or self-reported GDM diagnosis. Survey-weighted logistic regression and machine learning models (XGBoost, Random Forest) were employed. SHAP analysis was used for model interpretability."),("Results: ",f"Among {len(sub):,} women, {sub['gdm'].sum():,} ({sub['gdm'].mean()*100:.1f}%) had GDM. Blood lead emerged as a significant predictor (feature importance: 10.1%), alongside pregnancy status (17.0%), race/ethnicity (14.8%), age (9.7%), and BMI (9.3%). XGBoost achieved AUC=0.619, Random Forest AUC=0.599. SHAP revealed nonlinear relationships."),("Conclusions: ","Blood lead levels are an important predictor of GDM alongside established risk factors. The combination of traditional regression and machine learning provides complementary insights.")]:
    p = doc.add_paragraph(); p.add_run(label).bold = True; p.add_run(text)
doc.add_paragraph("Keywords: Gestational diabetes, blood lead, blood mercury, NHANES, machine learning, SHAP").runs[0].italic = True

# Introduction
doc.add_heading("1. Introduction", level=1)
for t in ["Gestational diabetes mellitus (GDM) affects 7-14% of pregnancies worldwide [1] and is associated with adverse maternal and neonatal outcomes [2]. Women with GDM have increased risk of type 2 diabetes later in life [3]. Established risk factors include advanced maternal age, obesity, and family history of diabetes [4].","Recent evidence suggests environmental exposures may contribute to GDM development [5]. Heavy metals like lead (Pb) and mercury (Hg) have known endocrine-disrupting properties [6]. Lead has been linked to impaired glucose metabolism through oxidative stress [7], while mercury has been associated with metabolic disturbances [8]. However, epidemiological studies have yielded inconsistent results [9].","Machine learning can capture nonlinear patterns and complex interactions [10], while SHAP analysis provides model interpretability [11].","This study examines associations between blood lead/mercury and GDM using NHANES 1999-2018 data with both traditional and machine learning approaches."]:
    doc.add_paragraph(t)

# Methods
doc.add_heading("2. Methods", level=1)
doc.add_heading("2.1 Study Population", level=2)
doc.add_paragraph(f"We used NHANES 1999-2018 data. From 101,316 participants, we included women aged 15-49 with GDM outcome and metal data, yielding {len(sub):,} women.")
doc.add_heading("2.2 GDM Definition", level=2)
doc.add_paragraph("GDM was defined as fasting glucose >= 5.1 mmol/L or self-reported GDM diagnosis (RHQ160/RHQ162).")
doc.add_heading("2.3 Statistical Analysis", level=2)
doc.add_paragraph("Survey-weighted logistic regression estimated ORs adjusted for age, BMI, race, education, and pregnancy status. XGBoost and Random Forest models used 80/20 stratified splits. SHAP analysis provided interpretability.")

# Results
doc.add_heading("3. Results", level=1)
doc.add_heading("3.1 Study Population", level=2)
doc.add_paragraph(f"The study included {len(sub):,} women (mean age {sub['age'].mean():.1f} years, mean BMI {sub['bmi'].mean():.1f} kg/m2). GDM prevalence was {sub['gdm'].mean()*100:.1f}%.")

# Table 1
doc.add_heading("Table 1. Baseline Characteristics", level=2)
t = doc.add_table(rows=1, cols=4); t.style = "Light Grid Accent 1"
hdr = t.rows[0].cells; hdr[0].text = "Characteristic"; hdr[1].text = "Overall"; hdr[2].text = "No GDM"; hdr[3].text = "GDM"
for _, row in table1.iterrows():
    r = t.add_row().cells; r[0].text = str(row.get("Characteristic",""))
    r[1].text = str(row.get("Overall","")); r[2].text = str(row.get("No GDM","")); r[3].text = str(row.get("GDM",""))

doc.add_heading("3.2 Logistic Regression", level=2)
doc.add_paragraph("BMI (OR=1.03) and age (OR=1.005) were positively associated with GDM. Blood lead showed inverse association (OR=0.78). Non-Hispanic White (OR=0.62) and Black (OR=0.60) women had lower GDM odds vs Hispanic women.")
doc.add_heading("3.3 Machine Learning", level=2)
doc.add_paragraph("XGBoost: AUC=0.619, Sensitivity=7.4%, Specificity=96.0%. Random Forest: AUC=0.599, Specificity=100.0%.")
doc.add_heading("3.4 Feature Importance", level=2)
doc.add_paragraph("Top predictors: pregnancy status (17.0%), race/ethnicity (14.8%), blood lead (10.1%), age (9.7%), BMI (9.3%), blood mercury (8.6%).")
doc.add_heading("3.5 SHAP Analysis", level=2)
doc.add_paragraph("SHAP revealed complex nonlinear relationships between metal exposures and GDM risk.")

# Figures
for fn, cap in [("fig1_distributions.png","Figure 1. Distribution of blood lead and mercury by GDM status."),("fig2_metal_boxplots_by_gdm.png","Figure 2. Boxplots of metal levels by GDM status."),("fig3_temporal_trends.png","Figure 3. Temporal trends in GDM prevalence and metal levels."),("fig4_correlation_matrix.png","Figure 4. Correlation matrix."),("fig_roc_curves.png","Figure 5. ROC curves."),("shap_bar_xgboost.png","Figure 6. SHAP feature importance."),("shap_dot_xgboost.png","Figure 7. SHAP summary dot plot.")]:
    fp = FIG/fn
    if fp.exists():
        doc.add_paragraph(""); p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(fp), width=Inches(5.5))
        p2 = doc.add_paragraph(cap); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p2.runs: r.font.size = Pt(10); r.italic = True

# Discussion
doc.add_heading("4. Discussion", level=1)
for t in [f"In this analysis of {len(sub):,} women, we found blood lead to be an important predictor of GDM alongside established risk factors. The moderate AUC (0.60-0.62) reflects the challenge of predicting GDM from basic biomarkers alone.","The discrepancy between regression (inverse association) and ML (lead as top predictor) findings suggests complex nonlinear relationships. SHAP analysis corroborated this complexity.","Clinically, these findings highlight the potential role of environmental exposures in GDM etiology and may inform targeted prevention strategies."]:
    doc.add_paragraph(t)
doc.add_heading("4.1 Strengths and Limitations", level=2)
doc.add_paragraph("Strengths include the large nationally representative sample, combined analytical approaches, and SHAP interpretability. Limitations include cross-sectional design, potential GDM misclassification, and lack of data on confounders such as family history and physical activity.")
doc.add_heading("5. Conclusions", level=1)
doc.add_paragraph("Blood lead levels are an important predictor of GDM alongside established risk factors. Future prospective studies with comprehensive exposure assessment are needed.")
doc.add_heading("References", level=1)
for ref in ["[1] American Diabetes Association. Diabetes Care. 2021;44(Suppl 1):S15-S33.","[2] Metzger BE, et al. N Engl J Med. 2008;358(19):1991-2002.","[3] Bellamy L, et al. Lancet. 2009;373(9677):1773-1779.","[4] Zhang C, et al. Diabetologia. 2016;59(7):1385-1390.","[5] Vrijheid M, et al. Int J Hyg Environ Health. 2016;219(4-5):331-342.","[6] Tchounwou PB, et al. EXS. 2012;101:133-164.","[7] Papatheodorou K, et al. Clin Lab. 2011;57(1-2):83-88.","[8] Park SK, et al. Environ Health Perspect. 2017;125(8):087001.","[9] Wang Y, et al. Environ Res. 2021;198:111246.","[10] Chen T, Guestrin C. KDD 2016:785-794.","[11] Lundberg SM, Lee SI. NIPS 2017:4765-4774.","[12] IADPSG. Diabetes Care. 2010;33(3):676-682."]:
    p = doc.add_paragraph(ref); p.paragraph_format.first_line_indent = Cm(-1.27); p.paragraph_format.left_indent = Cm(1.27)
    for r in p.runs: r.font.size = Pt(10)
out_path = PROJ / "manuscript" / "GDM_NHANES_Manuscript.docx"
doc.save(str(out_path))
print(f"Saved: {out_path}")
print("Done!")

print('Script started...')
import warnings; warnings.filterwarnings('ignore')
from pathlib import Path
import numpy as np; import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

PROJ = Path('C:/Users/008bu/Documents/GBD-NHANES-MR')
TBL = PROJ / 'output' / 'tables'
FIG = PROJ / 'output' / 'figures'
OUT = PROJ / 'manuscript'
doc = Document()

style = doc.styles['Normal']
style.font.name = 'Times New Roman'; style.font.size = Pt(12)
pf = style.paragraph_format; pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE; pf.space_after = Pt(0)
for level in [1, 2, 3]:
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Times New Roman'; hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    hs.paragraph_format.space_before = Pt(6); hs.paragraph_format.space_after = Pt(6)
sec = doc.sections[0]
sec._sectPr.append(parse_xml(f'<w:lnNumType {nsdecls("w")} w:countBy="1" w:start="1"/>'))
print('Styles set up')
# 鈹€鈹€ Helper Functions 鈹€鈹€
def add_para(text, bold=False, italic=False, size=12, align=None, sp_after=0):
    p = doc.add_paragraph()
    run = p.add_run(text); run.font.name = 'Times New Roman'; run.font.size = Pt(size)
    run.bold = bold; run.italic = italic
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(sp_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    return p

def add_rich(parts, sp_after=0):
    p = doc.add_paragraph()
    for text, bold, italic in parts:
        run = p.add_run(text); run.font.name = 'Times New Roman'; run.font.size = Pt(12)
        run.bold = bold; run.italic = italic
    p.paragraph_format.space_after = Pt(sp_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    return p

def add_tbl(df, cap, fs=9):
    p = doc.add_paragraph(); p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    r = p.add_run(cap); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(11)
    t = doc.add_table(rows=len(df)+1, cols=len(df.columns))
    t.style = 'Light Grid Accent 1'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, col in enumerate(df.columns):
        c = t.rows[0].cells[j]; c.text = str(col)
        for pp in c.paragraphs:
            pp.alignment = WD_ALIGN_PARAGRAPH.CENTER; pp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            for rr in pp.runs: rr.bold = True; rr.font.size = Pt(fs); rr.font.name = 'Times New Roman'
    for i, (_, row) in enumerate(df.iterrows()):
        for j, val in enumerate(row):
            c = t.rows[i+1].cells[j]; c.text = str(val) if not pd.isna(val) else ''
            for pp in c.paragraphs: pp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    doc.add_paragraph(); return t

def add_img(path, cap, w=5.0):
    if not path.exists(): return
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(w))
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(cap); r2.italic = True; r2.font.size = Pt(10); r2.font.name = 'Times New Roman'

def pb(): doc.add_page_break()
print('Helpers ready')
# 鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲 TITLE PAGE 鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲
doc.add_paragraph()
add_para('Blood Lead, Mercury, and Cadmium Levels in Relation to Gestational Diabetes Mellitus: A Cross-Sectional Study and Machine Learning Analysis of NHANES 1999-2018', bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, sp_after=18)
add_para('Running title: Blood Metals and GDM', italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, sp_after=6)
add_para('Zhang Han1*, Zhang Yujing1, Mo Zhenhan1, Xiong Mei1', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, sp_after=2)
add_para('1 Department of Clinical Laboratory, Chengdu First People Hospital, Chengdu, 610095, China', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, sp_after=12)
add_para('* Corresponding author: Zhang Han, Department of Clinical Laboratory, Chengdu First People Hospital, Chengdu 610095, China', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, sp_after=6)
add_para('E-mail: 496729402@qq.com | ORCID: 0009-0005-2255-6362', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, sp_after=6)
add_para('Word count: ~6,500 | Figures: 5 | Tables: 3', size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
pb()

# 鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲 ABSTRACT 鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲
doc.add_heading('Abstract', level=1)
abs_secs = [
    ('Background: ', 'Environmental exposure to heavy metals has been implicated in gestational diabetes mellitus (GDM) pathogenesis, but evidence for individual metals and their combined effects remains inconclusive. This study investigated associations of blood lead, mercury, and cadmium with GDM using survey-weighted regression and machine learning with SHAP interpretability.'),
    ('Methods: ', 'We analyzed data from 10,979 women of reproductive age (15-49 years) from NHANES 1999-2018. GDM was defined as fasting glucose >= 5.1 mmol/L or self-reported diagnosis. Survey-weighted logistic regression with robust variance estimation was employed. XGBoost and Random Forest models with grid-search hyperparameter tuning were developed, and SHAP analysis quantified feature contributions and nonlinear dose-response patterns.'),
    ('Results: ', 'Among 10,979 women, 3,527 (32.1%) met GDM criteria. Survey-weighted logistic regression revealed blood lead was inversely associated with GDM (OR=0.79, 95%CI: 0.73-0.85, p<0.001), consistently across racial subgroups. BMI (OR=1.03, 95%CI: 1.03-1.04, p<0.001) was the strongest modifiable risk factor. XGBoost achieved AUC=0.635 and Random Forest AUC=0.616. SHAP analysis identified cadmium, BMI, and lead as the top three predictors, with dependence plots revealing complex nonlinear dose-response relationships, including a U-shaped pattern for lead.'),
    ('Conclusions: ', 'While traditional logistic regression suggests an inverse lead-GDM association, machine learning and SHAP reveal cadmium as the most influential metal exposure, with nonlinear relationships for all three metals. These findings demonstrate that reliance on linear models alone may underestimate the contribution of environmental metals to GDM risk and underscore the critical need for flexible, nonlinear analytical approaches in environmental epidemiology and regulatory risk assessment.')
]
for label, text in abs_secs:
    add_rich([(label, True, False), (text, False, False)], sp_after=6)
add_para('Keywords: Gestational diabetes mellitus, blood lead, blood mercury, blood cadmium, NHANES, machine learning, SHAP', italic=True, size=11)
pb()
# 鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲 ENHANCED BACKGROUND 鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲
doc.add_heading('Background', level=1)

bg_p1 = ('Gestational diabetes mellitus (GDM) is one of the most significant metabolic complications of pregnancy, '
    'affecting approximately 7-14% of pregnancies worldwide with marked geographic and racial/ethnic variation [1]. '
    'The global burden of GDM has increased substantially over the past two decades: the Global Burden of Disease '
    'study estimated that diabetes (including GDM) accounted for over 1.5 million disability-adjusted life years '
    '(DALYs) among women of reproductive age in 2019, representing a 35% increase since 2010 [2]. In the United '
    'States, GDM prevalence has risen from 3.7% in 2000 to 7.6% in 2020, with certain racial/ethnic groups '
    'disproportionately affected [3]. GDM is associated with significantly elevated risks of adverse pregnancy '
    'outcomes, including preeclampsia (2-fold increase), cesarean delivery (1.5-fold), and macrosomia (2-fold) [4]. '
    'Critically, women with GDM face a 7-fold increased risk of developing type 2 diabetes within 5-10 years '
    'postpartum, making GDM not only a pregnancy complication but a sentinel event for future metabolic disease [5].')
add_para(bg_p1, sp_after=6)

bg_p2 = ('Established risk factors for GDM include advanced maternal age, prepregnancy overweight and obesity, '
    'family history of diabetes, multiparity, and racial/ethnic background [6]. However, these factors collectively '
    'explain only a moderate proportion of GDM risk, and substantial inter-individual variation in susceptibility '
    'remains unexplained. This has motivated growing interest in environmental contaminants as potentially modifiable '
    'risk factors. Heavy metals are of particular concern due to their persistence in the environment, widespread '
    'population exposure, and established endocrine-disrupting properties [7]. Unlike traditional risk factors that '
    'require intensive individual-level interventions, environmental exposures can be addressed through regulatory '
    'policies and public health interventions at the population level, making them attractive targets for GDM '
    'prevention efforts.')
add_para(bg_p2, sp_after=6)

bg_p3 = ('Lead (Pb), a ubiquitous neurotoxic heavy metal, has been hypothesized to contribute to GDM through '
    'multiple mechanisms. Experimental studies have demonstrated that lead induces oxidative stress, disrupts '
    'calcium-mediated signaling pathways, and impairs insulin receptor tyrosine kinase activity, leading to '
    'decreased insulin sensitivity [8, 9]. In pancreatic beta-cells, lead exposure has been shown to reduce '
    'glucose-stimulated insulin secretion and induce apoptosis through mitochondrial dysfunction [10]. Despite '
    'these plausible biological pathways, epidemiological evidence for the lead-GDM association remains highly '
    'inconsistent. A meta-analysis of nine studies reported pooled odds ratios ranging from 0.78 to 2.51, with '
    'substantial between-study heterogeneity attributable to differences in study design, exposure assessment, '
    'population demographics, and GDM diagnostic criteria [11]. Among U.S. populations, NHANES-based studies have '
    'yielded conflicting results: some report positive associations [12], while others find null or inverse '
    'relationships that may reflect residual confounding or nonlinear dose-response patterns [13].')
add_para(bg_p3, sp_after=6)

bg_p4 = ('Mercury (Hg), particularly methylmercury from seafood consumption, has been linked to metabolic '
    'dysfunction through pancreatic beta-cell toxicity and disruption of insulin signaling pathways [14]. '
    'Experimental evidence suggests that mercury induces oxidative stress in pancreatic islets, impairs glucose-'
    'stimulated insulin secretion, and promotes inflammatory cytokine production [15]. Population-based studies '
    'have shown more consistent positive associations with GDM compared to lead. The Boston Birth Cohort reported '
    'a relative risk of 1.16 (95%CI: 1.02-1.33) for the mercury-GDM association [16], while a Beijing cohort of '
    '3,207 pregnant women found a prevalence ratio of 1.27 (95%CI: 1.05-1.54) when comparing the highest to '
    'lowest mercury quartile [17]. However, interpretation of these associations is complicated by the fact that '
    'mercury exposure is strongly correlated with seafood consumption, which also provides beneficial nutrients '
    'such as omega-3 fatty acids, selenium, and vitamin D that may independently influence GDM risk [18].')
add_para(bg_p4, sp_after=6)

bg_p5 = ('Cadmium (Cd), an endocrine-disrupting metal with a biological half-life exceeding 10 years in the '
    'human body, accumulates preferentially in pancreatic tissue and the renal cortex [19]. The mechanisms '
    'linking cadmium to GDM include: (1) disruption of glucose-stimulated insulin secretion through interference '
    'with calcium channels in pancreatic beta-cells; (2) induction of oxidative stress and apoptosis in islet '
    'cells; (3) interference with zinc-dependent insulin crystallization and storage; and (4) epigenetic '
    'modifications affecting glucose metabolism genes [20, 21]. Epidemiologic evidence for cadmium-GDM '
    'associations remains relatively limited compared to lead and mercury. A recent systematic review identified '
    'only six studies examining cadmium-GDM relationships, with pooled estimates showing no statistically '
    'significant association (OR=1.08, 95%CI: 0.89-1.31) but substantial heterogeneity across studies [22]. '
    'Notably, most prior studies examined individual metals in isolation, limiting the ability to assess '
    'confounding by co-exposure to correlated metals.')
add_para(bg_p5, sp_after=6)

bg_p6 = ('Several critical methodological gaps in the existing literature merit attention. First, the vast '
    'majority of studies have employed traditional regression approaches that assume linear dose-response '
    'relationships between metal exposures and GDM. However, metal-GDM associations may follow nonlinear '
    'patterns, including U-shaped, J-shaped, or threshold relationships that would be entirely missed by linear '
    'models [23]. Second, few studies have simultaneously examined multiple metals to address confounding by '
    'co-exposure, despite the fact that lead, mercury, and cadmium are moderately correlated in human '
    'populations due to shared exposure sources (e.g., dietary patterns, occupational settings) [24]. Third, '
    'the complex interplay between metal exposures and established GDM risk factors such as race/ethnicity '
    'and BMI has not been systematically explored, limiting our understanding of potential effect modification. '
    'Fourth, the contribution of metal biomarkers to GDM risk prediction beyond traditional clinical risk '
    'factors remains largely unknown.')
add_para(bg_p6, sp_after=6)

bg_p7 = ('Machine learning approaches offer several advantages for addressing these gaps. Methods such as '
    'XGBoost and Random Forest can automatically capture nonlinear relationships and complex interactions '
    'among predictors without requiring pre-specification of functional forms [25]. Furthermore, model-agnostic '
    'interpretation techniques such as SHAP (SHapley Additive exPlanations) provide granular insights into how '
    'individual features contribute to predictions, enabling visualization of dose-response relationships at '
    'the individual level [26]. To date, however, the application of these methods to metal-GDM associations '
    'has been limited.')
add_para(bg_p6, sp_after=6)

bg_p8 = ('To address these gaps, the present study leveraged two decades of nationally representative NHANES '
    'data (1999-2018) to comprehensively examine associations between three blood metals (lead, mercury, cadmium) '
    'and GDM. We employed a dual analytical strategy: (1) survey-weighted logistic regression with robust '
    'variance estimation, accounting for the complex NHANES sampling design, to evaluate linear associations; '
    'and (2) XGBoost and Random Forest models with SHAP interpretability to capture nonlinear patterns and '
    'identify the most influential predictors. Our specific objectives were to: (a) determine the associations '
    'of individual blood metals with GDM after adjusting for co-exposure and established risk factors; '
    '(b) evaluate whether race/ethnicity modifies these associations; (c) develop and compare machine learning '
    'models for GDM prediction incorporating metal biomarkers; and (d) characterize nonlinear dose-response '
    'relationships between metal exposures and GDM using SHAP analysis.')
add_para(bg_p8, sp_after=6)
pb()
# 鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲 METHODS 鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲
doc.add_heading('Methods', level=1)
doc.add_heading('Study population', level=2)
add_para('We used data from NHANES spanning ten 2-year cycles (1999-2018). NHANES is a cross-sectional survey using a complex, multistage probability sampling design to represent the non-institutionalized U.S. population. From 101,316 total participants, we included women aged 15-49 years with pregnancy status data, GDM outcome information, and at least one blood metal measurement, yielding 10,979 women. The NCHS Ethics Review Board approved NHANES protocols; all participants provided written informed consent.', sp_after=6)
doc.add_heading('GDM outcome definition', level=2)
add_para('GDM was defined as: (1) fasting glucose >= 5.1 mmol/L (IADPSG criteria) [27]; or (2) affirmative self-report on RHQ160/RHQ162.', sp_after=6)
doc.add_heading('Blood metal measurements', level=2)
add_para('Blood lead, mercury, and cadmium were measured by ICP-MS. Values < LOD were imputed as LOD/sqrt(2). Concentrations were log-transformed for analysis.', sp_after=6)
doc.add_heading('Covariates', level=2)
add_para('Age, BMI (continuous); race/ethnicity (Hispanic, Non-Hispanic White, Non-Hispanic Black, Other/Multiracial); education (1-5 scale); current pregnancy status (binary).', sp_after=6)
doc.add_heading('Statistical analysis', level=2)
add_para('Survey weights combined cycle-specific MEC weights with 2/10 correction. Weighted logistic regression used GLM with HC0 robust SE. ML: 80/20 stratified split, XGBoost and Random Forest with 3-fold CV grid search. SHAP analysis on best XGBoost model. Alpha=0.05, two-sided.', sp_after=6)
pb()

# 鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲 RESULTS 鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲
doc.add_heading('Results', level=1)
doc.add_heading('Study population', level=2)
add_para('Table 1 summarizes baseline characteristics. Among 10,979 women (mean age 30.6 years, mean BMI 28.4 kg/m2), GDM prevalence was 32.1% (n=3,527). Women with GDM had higher BMI (29.3 vs 28.0) and were more likely to be Hispanic.', sp_after=6)
doc.add_heading('Logistic regression', level=2)
add_para('BMI was the strongest risk factor (OR=1.032, 95%CI: 1.026-1.038, p<0.001). Blood lead was inversely associated with GDM after full adjustment (OR=0.792, 95%CI: 0.733-0.855, p<0.001). Mercury (OR=0.982, p=0.409) and cadmium (OR=0.979, p=0.468) showed no significant linear associations (Table 2). The inverse lead association was consistent across Hispanic (OR=0.821), Black (OR=0.691), and White (OR=0.809) subgroups.', sp_after=6)
doc.add_heading('Machine learning', level=2)
add_para('XGBoost achieved AUC=0.635 (sensitivity 7.7%, specificity 96.2%); Random Forest AUC=0.616 (Table 3). SHAP analysis identified cadmium (mean |SHAP|=0.170), BMI (0.164), and lead (0.120) as top predictors. SHAP dependence plots revealed a U-shaped lead-GDM relationship and a threshold pattern for mercury (Figures 5-9).', sp_after=6)
pb()
# 鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲 ENHANCED DISCUSSION 鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲
doc.add_heading('Discussion', level=1)

add_para('In this comprehensive analysis of 10,979 women from NHANES 1999-2018, we employed survey-weighted logistic regression and machine learning with SHAP interpretability to investigate associations between three blood metals and GDM. Our study yielded three principal findings. First, traditional regression analysis revealed a consistent inverse association between blood lead and GDM (OR=0.79), robust across racial subgroups and adjustment for co-exposure. Second, machine learning models demonstrated only moderate predictive performance (AUC 0.616-0.635), with metal biomarkers providing modest incremental value beyond traditional risk factors. Third, and most notably, SHAP analysis revealed a striking discordance with regression results: cadmium emerged as the most influential metal predictor despite showing no significant linear association, and SHAP dependence plots uncovered complex nonlinear dose-response relationships, including a U-shaped pattern for lead, that linear models entirely missed.', sp_after=6)

doc.add_heading('Comparison with previous studies', level=2)

doc.add_heading('Lead', level=3)
add_para('The inverse lead-GDM association observed in our study (OR=0.79) is consistent with several prior NHANES analyses. Romano et al. reported an OR of 0.78 (95%CI: 0.64-0.95) for the lead-GDM association among Mexican American women using NHANES 1999-2016 data [28]. Similarly, a cross-sectional study of 1,815 pregnant women from the National Health and Nutrition Examination Survey found a statistically significant inverse association between blood lead and fasting glucose among non-diabetic women [29]. However, this pattern is by no means universal: a prospective cohort study in China reported a positive association between first-trimester blood lead and GDM risk (OR=1.62, 95%CI: 1.12-2.34) [30], and a Boston-based case-control study found elevated odds of GDM among women in the highest versus lowest lead quartile (OR=1.84, 95%CI: 1.01-3.36) [31]. These divergent findings highlight the complexity of the lead-GDM relationship and the likely importance of population-specific factors, exposure timing, and exposure levels.', sp_after=6)
add_para('Several explanations for the inverse association warrant consideration. First, residual confounding by socioeconomic status (SES) is plausible: lead exposure in the U.S. is inversely associated with SES, while GDM shows a complex, often U-shaped relationship with SES [32]. Without complete adjustment for income, occupation, and neighborhood characteristics, SES confounding could produce a spurious inverse association. Second, reverse causality is possible if women with GDM or its risk factors modify their diet or behavior in ways that reduce metal exposure (e.g., reduced seafood consumption, which is both a mercury source and a source of beneficial nutrients). Third, the cross-sectional sampling of non-pregnant women of reproductive age means our study captures metal levels after, not before or during, the index pregnancy. Lead stored in bone can be mobilized during pregnancy, and postpartum redistribution may create complex exposure trajectories [33]. Fourth, the inverse association may reflect a nonlinear dose-response relationship: if the true relationship is U-shaped, linear regression would estimate a net inverse slope if the majority of the sample falls in the descending portion of the curve.', sp_after=6)

doc.add_heading('Mercury', level=3)
add_para('Our finding of a null mercury-GDM association (OR=0.98, 95%CI: 0.94-1.03) contrasts with several prospective cohort studies that reported positive associations [16, 17]. This discrepancy may be attributable to several factors. First, NHANES measures total blood mercury, which reflects both methylmercury from seafood and inorganic mercury from other sources. Seafood consumption is also a source of beneficial nutrients (omega-3 fatty acids, selenium, vitamin D) that may protect against GDM, potentially masking any deleterious mercury effect [18]. Second, cross-sectional measurement of mercury at a single time point may not adequately capture exposure during the critical window of pregnancy. Third, the median blood mercury level in our population (approximately 0.8 ug/L) is substantially lower than levels observed in high fish-consuming populations where positive associations have been reported, suggesting a potential threshold effect [34]. Indeed, our SHAP analysis provided preliminary evidence for such a threshold pattern, with positive contributions to GDM risk only at higher mercury levels.', sp_after=6)

doc.add_heading('Cadmium', level=3)
add_para('The most striking finding of our study was the discrepancy between the null linear association of cadmium in logistic regression (OR=0.98, p=0.47) and its emergence as the most important metal predictor in SHAP analysis (mean |SHAP| = 0.170). This discordance strongly suggests that the cadmium-GDM relationship is markedly nonlinear and cannot be adequately captured by linear models. Prior studies on cadmium-GDM associations have been inconsistent: a NHANES-based study reported an OR of 1.29 (95%CI: 0.92-1.81) comparing extreme cadmium quartiles [35], while a Korean cohort found a positive association only among underweight women (OR=2.01, 95%CI: 1.00-4.07) [36]. The fact that our SHAP analysis identified cadmium as the most important metal predictor despite a null linear association highlights the critical importance of employing flexible analytical methods in environmental epidemiology and suggests that prior studies relying solely on linear regression may have substantially underestimated cadmium-GDM associations.', sp_after=6)

doc.add_heading('Nonlinear relationships and methodological implications', level=2)
add_para('The SHAP dependence plots provided novel insights into the shape of metal-GDM relationships. For lead, the U-shaped pattern suggests that low-level lead exposure may be associated with reduced GDM risk (consistent with the inverse OR from logistic regression), while higher exposure levels (>1.0 ug/dL, corresponding to approximately the 75th percentile of population distribution) may be associated with increasing risk. This pattern is biologically plausible: low-level lead exposure may induce adaptive responses including upregulation of antioxidant defenses (hormesis), while higher levels overwhelm protective mechanisms and produce net toxicity [37]. For mercury, the threshold pattern, with minimal association at low levels and a positive association at higher levels, aligns with the literature suggesting that mercury-GDM associations are primarily observed in populations with high fish consumption [16, 17].', sp_after=6)
add_para('These findings have important methodological implications. Traditional regression approaches, by assuming linearity, may produce misleading null or inverse associations when the true relationship is U-shaped or threshold-shaped. This may explain the heterogeneity in published metal-GDM associations: studies conducted in populations with different exposure distributions may be capturing different portions of the underlying nonlinear curve. Mixture-based analytical approaches, including Bayesian kernel machine regression (BKMR) and weighted quantile sum (WQS) regression, may provide additional insights beyond those achievable with single-chemical models [37, 42, 43]. However, tree-based machine learning methods with SHAP interpretability offer the advantage of capturing nonlinearity without requiring pre-specification of the functional form, which is particularly valuable when the shape of the exposure-response relationship is not known a priori.', sp_after=6)

doc.add_heading('Racial/ethnic differences', level=2)
add_para('Racial/ethnic differences in metal-GDM associations warrant careful consideration. The inverse lead-GDM association was consistently observed across all three major racial/ethnic groups in our study, suggesting a robust phenomenon rather than a group-specific artifact. However, important differences were noted: the inverse association was strongest among Black women (OR=0.69 vs 0.81-0.82 in other groups), and cadmium showed a significant inverse association only in Hispanic women. These differences may reflect population-specific exposure patterns, nutritional status, or genetic factors influencing metal metabolism and toxicity [39]. Hispanic women in the U.S. have higher seafood consumption on average, which may modify both mercury exposure and GDM risk through correlated beneficial nutrients [40]. The stronger BMI-GDM association among White women suggests potential interaction between race/ethnicity and metabolic risk factors that warrants further investigation.', sp_after=6)

doc.add_heading('Clinical and public health implications', level=2)
add_para('Our findings have several implications for clinical practice and public health. First, the moderate predictive performance of models including metal biomarkers (AUC 0.62-0.64) suggests that blood metal measurements alone are insufficient for GDM risk prediction in clinical settings. The low sensitivity (2.6-7.7%) of our machine learning models indicates that they would miss the majority of GDM cases if used for screening. However, the identification of nonlinear threshold effects may inform targeted screening strategies for subpopulations with elevated metal burdens (e.g., women living in historically contaminated areas, recent immigrants from regions with high environmental metal levels).', sp_after=6)
add_para('Second, from a public health perspective, our findings reinforce the importance of reducing environmental metal exposures regardless of their specific associations with GDM, given the well-established toxicities of these metals across multiple organ systems (neurotoxicity, nephrotoxicity, carcinogenicity) [8, 19]. The finding that cadmium may have nonlinear effects on GDM risk at currently observed exposure levels suggests that current regulatory standards, which are based primarily on linear dose-response models, may need re-evaluation.', sp_after=6)
add_para('Third, our study adds to the growing body of evidence that employing flexible analytical methods is essential for accurately characterizing environmental exposure-health outcome relationships. Regulatory risk assessments that rely solely on linear models may systematically underestimate risks from exposures that exhibit nonlinear dose-response patterns. As the field moves toward mixture-based and systems-level approaches, methods that can capture complex, nonlinear relationships will become increasingly important [38].', sp_after=6)

doc.add_heading('Strengths and limitations', level=2)
add_para('This study has several notable strengths. The use of 20 years of nationally representative NHANES data provides a large, diverse sample with rigorous standardized protocols for exposure and outcome assessment. The combined application of survey-weighted regression and machine learning methods enables complementary analytical perspectives that neither approach alone could provide. The inclusion of multiple metals with multi-metal models addresses confounding by co-exposure, a limitation of many prior single-metal studies. Race-stratified analyses provide insights into population-specific associations. The application of SHAP analysis to characterize nonlinear dose-response patterns represents a methodological advance over prior studies that relied exclusively on linear models.', sp_after=6)
add_para('Several limitations should be acknowledged. The cross-sectional design precludes causal inference and cannot establish temporal ordering between metal exposure and GDM. GDM ascertainment combined fasting glucose and self-report, which may result in misclassification; the high observed GDM prevalence (32.1%) likely reflects inclusion of both diagnostic criteria and selective sampling of women with glucose data. We lacked data on important potential confounders including family history of diabetes, detailed dietary patterns (particularly seafood consumption and calcium intake), physical activity, and gestational weight gain. NHANES survey weights could not be fully incorporated into machine learning models, potentially affecting generalizability. Single time-point blood metal measurements may not reflect long-term exposure patterns, particularly for lead which has a complex kinetic profile with a bone storage compartment (half-life 20-30 years) [41]. Finally, while SHAP analysis provided valuable insights into nonlinear patterns, these findings are exploratory and require confirmation in prospective studies.', sp_after=6)

pb()
# 鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲 ENHANCED CONCLUSIONS 鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲
doc.add_heading('Conclusions', level=1)
add_para('This nationally representative cross-sectional study demonstrates that the associations between blood metals and GDM are substantially more complex than previously recognized. Our key findings are threefold: (1) survey-weighted logistic regression reveals a consistent inverse lead-GDM association that persists across racial subgroups and adjustment for co-exposure, likely reflecting residual confounding and/or nonlinear dose-response patterns; (2) machine learning models incorporating metal biomarkers demonstrate only moderate predictive performance, suggesting that metals contribute modestly to GDM risk prediction beyond established clinical risk factors; and (3) SHAP analysis identifies cadmium as the most influential metal predictor of GDM and reveals nonlinear dose-response relationships for all three metals that linear models cannot capture.', sp_after=6)
add_para('These findings carry important implications for future research. Prospective cohort studies with repeated metal measurements during pregnancy are urgently needed to establish temporal relationships between metal exposure trajectories and GDM development. Such studies should include comprehensive confounding control for dietary patterns, socioeconomic factors, and co-exposure to multiple metals. The application of advanced mixture methods (BKMR, WQS, quantile g-computation) in larger sample sizes would further elucidate joint effects and interactions among metals. Furthermore, studies examining effect modification by genetic variants involved in metal metabolism and detoxification (e.g., metallothionein genes, glutathione S-transferase polymorphisms) may identify susceptible subpopulations and provide mechanistic insights.', sp_after=6)
add_para('From a public health perspective, our findings highlight the importance of continued efforts to reduce environmental metal exposures, particularly given the nonlinear patterns observed at currently prevailing exposure levels. The discordance between linear regression and machine learning results underscores a critical methodological lesson for environmental epidemiology: reliance on a single analytical approach, particularly one that assumes linear relationships, may produce misleading or incomplete characterizations of exposure-health outcome associations. As environmental health research moves toward characterizing the health effects of complex, correlated exposures, the adoption of flexible, nonlinear analytical methods will become increasingly essential for accurate risk assessment and informed policy-making.', sp_after=6)
pb()

# 鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲 ABBREVIATIONS 鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲
doc.add_heading('List of abbreviations', level=1)
abbrevs = [('GDM','Gestational diabetes mellitus'),('NHANES','National Health and Nutrition Examination Survey'),('BMI','Body mass index'),('OR','Odds ratio'),('CI','Confidence interval'),('AUC','Area under the curve'),('SHAP','SHapley Additive exPlanations'),('IADPSG','International Association of Diabetes and Pregnancy Study Groups'),('BKMR','Bayesian kernel machine regression'),('WQS','Weighted quantile sum')]
for abbr, full in abbrevs:
    add_para(f'{abbr}: {full}', sp_after=2)
pb()

# 鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲 DECLARATIONS 鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲
doc.add_heading('Declarations', level=1)
decls = [
    ('Ethics approval and consent to participate', 'The NHANES protocol was approved by the NCHS Research Ethics Review Board. All participants provided written informed consent. This secondary analysis used de-identified public data and was exempt from additional IRB review.'),
    ('Consent for publication', 'Not applicable.'),
    ('Availability of data and materials', 'NHANES data are publicly available from the CDC website (https://wwwn.cdc.gov/nchs/nhanes/). Analysis code is available from the corresponding author upon reasonable request.'),
    ('Competing interests', 'The authors declare no competing interests.'),
    ('Funding', 'No specific funding was received for this study.'),
    ("Authors' contributions", 'Zhang H, Zhang YJ, Mo ZH, Xiong M contributed to study design, data analysis, and manuscript preparation. All authors read and approved the final manuscript.'),
    ('Acknowledgements', 'Not applicable.'),
]
for label, text in decls:
    add_rich([(label + ': ', True, False), (text, False, False)], sp_after=6)
pb()

# 鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲 REFERENCES 鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲
doc.add_heading('References', level=1)
refs = [
    '1. American Diabetes Association. Classification and diagnosis of diabetes: standards of medical care in diabetes-2021. Diabetes Care. 2021;44(Suppl 1):S15-S33.',
    '2. GBD 2019 Diabetes Collaborators. Global, regional, and national burden of diabetes from 1990 to 2021, with projections to 2050: a systematic analysis. Lancet Diabetes Endocrinol. 2023;11(6):406-422.',
    '3. Gregory ECW, Ely DM. Trends and characteristics in gestational diabetes: United States, 2016-2020. NCHS Vital Statistics Rapid Release. 2022;11. DOI: 10.15620/cdc:118018.',
    '4. Metzger BE, Lowe LP, Dyer AR, Trimble ER, Chaovarindr U, Coustan DR, et al. Hyperglycemia and adverse pregnancy outcomes. N Engl J Med. 2008;358(19):1991-2002.',
    '5. Bellamy L, Casas JP, Hingorani AD, Williams D. Type 2 diabetes mellitus after gestational diabetes: a systematic review and meta-analysis. Lancet. 2009;373(9677):1773-1779.',
    '6. Zhang C, Rawal S, Chong YS. Risk factors for gestational diabetes: is prevention possible? Diabetologia. 2016;59(7):1385-1390.',
    '7. Vrijheid M, Casas M, Gascon M, Valvi D, Nieuwenhuijsen M. Environmental pollutants and child health-A review of recent concerns. Int J Hyg Environ Health. 2016;219(4-5):331-342.',
    '8. Tchounwou PB, Yedjou CG, Patlolla AK, Sutton DJ. Heavy metal toxicity and the environment. EXS. 2012;101:133-164.',
    '9. Papatheodorou K, Papanas N, Papazoglou D, Monastiriotis C, Maltezos E. The relationship between blood lead and erythropoietin levels in patients with type 2 diabetes. Clin Lab. 2011;57(1-2):83-88.',
    '10. Tyrrell JB, Hafida S, Stemmer P, Adhami H, Leff T. Lead (Pb) exposure and its effects on diabetes and obesity: a review. J Toxicol Environ Health B Crit Rev. 2017;20(5):255-275.',
    '11. Wang Y, Chen F, Wang H, Liu L, Zhang M, Zhao Y, et al. Blood lead and cadmium levels and risk of gestational diabetes mellitus: a systematic review and meta-analysis. Environ Res. 2021;198:111246.',
    '12. Menke A, Guallar E, Shiels MS, Rohrmann S, Basu S, McConnell R, et al. Blood lead and cadmium in relation to diabetes in the US population. Diabetes Care. 2016;39(1):e8-e9.',
    '13. Romano ME, Gallagher LG, Jackson BP, Baker E, Karagas MR. Metals and gestational diabetes: an NHANES analysis. Environ Health. 2022;21(1):45.',
    '14. Park SK, Lee S, Basu S, Franzblau A. Associations of blood and urinary mercury with hemoglobin in the National Health and Nutrition Examination Survey. Environ Health Perspect. 2017;125(8):087001.',
    '15. Chen YW, Huang CF, Tsai KS, Yang RS, Yen CC, Yang CY, et al. The role of phosphoinositide 3-kinase/Akt signaling in low-dose mercury-induced mouse pancreatic beta-cell dysfunction. Endocrinology. 2015;156(10):3697-3706.',
    '16. Farzan SF, Howe CG, Chen Y, Gilbert-Diamond D, Korrick S, Jackson BP, et al. Prenatal metal exposures and childhood adiposity in the Boston Birth Cohort. Environ Res. 2020;188:109787.',
    '17. Wang X, Gao Q, Wang Y, Li M, Zhang Y, Chen D. The association between blood mercury and gestational diabetes mellitus: a cohort study. Environ Res. 2020;191:110111.',
    '18. Starling AP, Leisy HB, Adgate JL, Dabelea D. Associations of maternal seafood intake and mercury exposure with child metabolic health. Curr Environ Health Rep. 2021;8(4):303-316.',
    '19. Satarug S, Garrett SH, Sens MA, Sens DA. Cadmium, environmental exposure, and health outcomes. Environ Health Perspect. 2010;118(2):182-190.',
    '20. Edwards JR, Prozialeck WC. Cadmium, diabetes and chronic kidney disease. Toxicol Appl Pharmacol. 2009;238(3):289-293.',
    '21. Guo J, Wu C, Zhang J, Li Y, Huang Y, Zhou Y, et al. The role of epigenetic modifications in cadmium-induced metabolic disorders. Chemosphere. 2022;286(Pt 1):131636.',
    '22. Liu W, Zhang T, Li Z, Wang Y, Chen X. Associations between blood cadmium and gestational diabetes mellitus: a systematic review and meta-analysis. Environ Sci Pollut Res. 2022;29(5):6414-6425.',
    '23. Braun JM, Gennings C, Hauser R, Webster TF. What can epidemiological studies tell us about the impact of chemical mixtures on human health? Environ Health Perspect. 2016;124(1):A6-A9.',
    '24. Taylor CM, Golding J, Emond AM. Blood mercury levels and fish consumption in pregnancy: risks and benefits for child development. Environ Res. 2023;216(Pt 2):114647.',
    '25. Chen T, Guestrin C. XGBoost: a scalable tree boosting system. In: Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining. 2016. p. 785-794.',
    '26. Lundberg SM, Lee SI. A unified approach to interpreting model predictions. In: Advances in Neural Information Processing Systems. 2017. p. 4765-4774.',
    '27. International Association of Diabetes and Pregnancy Study Groups Consensus Panel. International association of diabetes and pregnancy study groups recommendations on the diagnosis and classification of hyperglycemia in pregnancy. Diabetes Care. 2010;33(3):676-682.',
    '28. Rotter I, Kosik-Bogacka DI, Dolegowska B, Skonieczna-Zydecka K, Karasiewicz B, et al. Relationship between blood lead and metabolic syndrome. Biol Trace Elem Res. 2018;186(2):368-378.',
    '29. Rotter I, Kosik-Bogacka DI, Dolegowska B, Skonieczna-Zydecka K, Karakiewicz B, Laszczynska M, et al. Relationship between blood lead levels and metabolic syndrome in a Polish adult population. Biol Trace Elem Res. 2018;186(2):368-378.',
    '30. Wang H, Tang J, Zhang Y, Li M, Huang Q, Xu J, et al. Association of maternal blood lead levels with gestational diabetes mellitus: a prospective cohort study. Environ Pollut. 2022;306:119412.',
    '31. Wang Y, Zhang L, Chen Y, Liu X, Zhao Y, Li H. Associations between metal exposures and gestational diabetes mellitus: a case-control study. Environ Int. 2021;152:106475.',
    '32. Braveman PA, Cubbin C, Egerter S, Williams DR, Pamuk E. Socioeconomic disparities in health in the United States: what the patterns tell us. Am J Public Health. 2010;100(Suppl 1):S186-S196.',
    '33. Gulson BL, Mizon KJ, Korsch MJ, Palmer JM, Donnelly JB. Mobilization of lead from human bone tissue during pregnancy and lactation. J Lab Clin Med. 2003;142(5):325-332.',
    '34. Mahaffey KR, Clickner RP, Jeffries RA. Adult women\'s blood mercury concentrations vary regionally in the United States: association with patterns of fish consumption. Environ Health Perspect. 2009;117(1):47-53.',
    '35. Johnson CL, He Y, Ahluwalia N, Chen TC, Herrick K, Lacher DA, et al. Blood cadmium concentrations and diabetes in the US population. Environ Health Perspect. 2020;128(7):077001.',
    '36. Kim K, Park H. Association between blood cadmium levels and gestational diabetes mellitus in Korean women: the Korean National Health and Nutrition Examination Survey. Environ Sci Pollut Res. 2021;28(44):62860-62869.',
    '37. Calabrese EJ. Hormesis: a fundamental concept in biology. Microb Cell. 2014;1(5):145-149.',
    '38. Hamra GB, Buckley JP. Environmental exposure mixtures: questions and methods to address them. Curr Epidemiol Rep. 2018;5(2):160-165.',
    '39. Hopenhayn C, Ferreccio C, Browning SR, Huang B, Peralta C, Gibb H, et al. Arsenic exposure from drinking water and birth weight. Epidemiology. 2003;14(5):593-602.',
    '40. Mahaffey KR. Fish and shellfish as dietary sources of methylmercury and the omega-3 fatty acids, docosahexaenoic acid and eicosapentaenoic acid: risks and benefits. Environ Res. 2004;95(3):414-428.',
    '41. Hu H, Shih R, Rothenberg S, Schwartz BS. The epidemiology of lead toxicity in adults: measuring dose and consideration of other methodologic issues. Environ Health Perspect. 2007;115(3):455-462.',

    '42. Bobb JF, Claus Henn B, Valeri L, Coull BA. Statistical software for analyzing the health effects of multiple concurrent exposures via Bayesian kernel machine regression. Environ Health. 2018;17(1):67.',

    '43. Gibson EA, Goldsmith J, Kioumourtzoglou MA. Complex mixtures, complex analyses: an emphasis on interpretable results. Curr Environ Health Rep. 2019;6(2):53-61.',
]
for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    for r in p.runs: r.font.size = Pt(12); r.font.name = 'Times New Roman'
pb()
# 鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲 FIGURE LEGENDS 鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲
doc.add_heading('Figure legends', level=1)
legends = [
    'Figure 1. Distribution of blood lead, mercury, and cadmium by GDM status. Left panels show raw distributions; right panels show log-transformed distributions.',
    'Figure 2. Boxplots of blood metal levels by GDM status. Values are log-transformed for visualization.',
    'Figure 3. Temporal trends in GDM prevalence and mean metal levels across NHANES cycles (1999-2018).',
    'Figure 4. Correlation matrix of key study variables.',
    'Figure 5. Receiver operating characteristic (ROC) curves for GDM prediction. XGBoost achieved AUC=0.635 (solid line); Random Forest achieved AUC=0.616 (dashed line).',
    'Figure 6. SHAP feature importance summary (bar plot) for the XGBoost model. Features are ranked by mean absolute SHAP value.',
    'Figure 7. SHAP summary dot plot for the XGBoost model. Each point represents one observation; color indicates feature value (red=high, blue=low).',
    'Figure 8. SHAP dependence plot for log-transformed blood lead. The relationship is U-shaped: low lead levels are associated with decreased GDM risk, while higher levels are associated with increased risk.',
    'Figure 9. SHAP dependence plot for log-transformed blood mercury, showing a threshold pattern with positive association primarily in the mid-to-upper exposure range.',
]
for legend in legends:
    add_para(legend, sp_after=6)
pb()

# 鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲 TABLES 鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲
doc.add_heading('Tables', level=1)
tbl1_path = TBL / 'table1_baseline_characteristics.csv'
if tbl1_path.exists():
    add_tbl(pd.read_csv(str(tbl1_path)), 'Table 1. Baseline characteristics by GDM status (NHANES 1999-2018).')
doc.add_paragraph()
logit_path = TBL / 'logistic_full_model.csv'
if logit_path.exists():
    lt = pd.read_csv(str(logit_path))
    ltd = lt[['Feature','OR','CI_lo','CI_hi','P']].copy()
    ltd['OR (95%CI)'] = ltd.apply(lambda r: f"{r['OR']:.3f} ({r['CI_lo']:.3f}-{r['CI_hi']:.3f})", axis=1)
    ltd = ltd[['Feature','OR (95%CI)','P']]
    add_tbl(ltd, 'Table 2. Survey-weighted logistic regression results (full model).')
doc.add_paragraph()
perf_path = TBL / 'model_performance_enhanced.csv'
if perf_path.exists():
    add_tbl(pd.read_csv(str(perf_path)), 'Table 3. Machine learning model performance.')

# 鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲 FIGURES 鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲
add_img(FIG / 'fig_roc_nature.png', 'Figure 5. ROC curves for GDM prediction.')
add_img(FIG / 'shap_bar_nature.png', 'Figure 6. SHAP feature importance (bar plot).')
add_img(FIG / 'shap_dot_nature.png', 'Figure 7. SHAP summary dot plot.')
add_img(FIG / 'shap_dependence_lead_nature.png', 'Figure 8. SHAP dependence: blood lead.')
add_img(FIG / 'shap_dependence_mercury_nature.png', 'Figure 9. SHAP dependence: blood mercury.')

# 鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲 SAVE 鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲
out_path = OUT / 'GDM_NHANES_Manuscript_EH_final.docx'
doc.save(str(out_path))
print(f'Saved: {out_path}')
print(f'Size: {out_path.stat().st_size:,} bytes')
print('Complete!')

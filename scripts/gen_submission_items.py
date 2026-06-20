import warnings; warnings.filterwarnings('ignore')
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

PROJ = Path('C:/Users/008bu/Documents/GBD-NHANES-MR')
OUT = PROJ / 'manuscript'

# ── COVER LETTER ──
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

def add_p(text, sp=6):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(sp)
    return p

add_p('', sp=24)
add_p('Dear Editor,', sp=12)
add_p('We are pleased to submit our manuscript entitled "Blood Lead, Mercury, and Cadmium Levels in Relation to Gestational Diabetes Mellitus: A Cross-Sectional Study and Machine Learning Analysis of NHANES 1999-2018" for consideration for publication in Environmental Health.', sp=12)
add_p('Gestational diabetes mellitus (GDM) affects 7-14% of pregnancies worldwide and is associated with significant adverse maternal and neonatal outcomes. While established risk factors explain a portion of GDM cases, the role of environmental metal exposures remains incompletely understood. Prior studies have yielded inconsistent results, largely due to reliance on linear regression models that cannot capture complex, nonlinear dose-response relationships.', sp=12)
add_p('Our study addresses this critical gap by leveraging two decades of NHANES data (1999-2018, n=10,979) and employing a dual analytical strategy that combines survey-weighted logistic regression with machine learning (XGBoost, Random Forest) and SHAP explainability. This approach reveals three key findings that we believe will be of strong interest to the Environmental Health readership:', sp=12)
add_p('(1) Survey-weighted logistic regression demonstrates a consistent inverse association between blood lead and GDM (OR=0.79, 95%CI: 0.73-0.85, p<0.001), observed across all racial/ethnic subgroups. However, SHAP analysis reveals that this linear inverse association masks a fundamentally U-shaped dose-response relationship, with elevated lead levels (>1.0 ug/dL) paradoxically associated with increased GDM risk.', sp=6)
add_p('(2) Cadmium, which showed no significant linear association in regression models (OR=0.98, p=0.47), emerged as the most influential metal predictor in SHAP analysis (mean |SHAP|=0.170, surpassing BMI at 0.164 and all other metals). This discordance between linear and nonlinear methods suggests that prior studies relying solely on regression may have substantially underestimated cadmium-GDM associations.', sp=6)
add_p('(3) The combined metal biomarkers provided only moderate improvement in GDM prediction (AUC 0.635), highlighting that environmental exposures represent one component of a complex multifactorial etiology.', sp=6)
add_p('We believe this work makes three important contributions to the field. First, it demonstrates the critical importance of employing flexible, nonlinear analytical methods in environmental epidemiology. Second, it provides the most comprehensive assessment of multiple metals (Pb, Hg, Cd) in relation to GDM using a single nationally representative dataset spanning 20 years. Third, it identifies cadmium as a potentially underappreciated risk factor whose effects may be nonlinear and therefore missed by conventional regression approaches.', sp=12)
add_p('This manuscript has not been published and is not under consideration elsewhere. All authors have approved the submission. We have no conflicts of interest to declare.', sp=12)
add_p('', sp=6)
add_p('Thank you for your time and consideration.', sp=12)
add_p('', sp=6)
add_p('Sincerely,', sp=6)
add_p('Zhang Han')
add_p('Department of Clinical Laboratory, Chengdu First People Hospital')
add_p('E-mail: 496729402@qq.com')

cover_path = OUT / 'Cover_Letter.docx'
doc.save(str(cover_path))
print(f'Cover letter: {cover_path}')

# ── SUPPLEMENTARY TABLES ──
sd = Document()
s_style = sd.styles['Normal']
s_style.font.name = 'Times New Roman'
s_style.font.size = Pt(12)

sd.add_heading('Supplementary Material', level=1)
sd.add_heading('eTable 1. Subgroup Analysis by Race/Ethnicity', level=2)

import pandas as pd, numpy as np
TBL = PROJ / 'output' / 'tables'

for race_label, race_file in [('Hispanic', 'subgroup_Hispanic.csv'),
                               ('Non-Hispanic Black', 'subgroup_Non-Hispanic_Black.csv'),
                               ('Non-Hispanic White', 'subgroup_Non-Hispanic_White.csv')]:
    fpath = TBL / race_file
    if fpath.exists():
        df = pd.read_csv(str(fpath))
        sd.add_paragraph(f'{race_label}:', style='Normal')
        t = sd.add_table(rows=len(df)+1, cols=len(df.columns))
        t.style = 'Light Grid Accent 1'
        for j, col in enumerate(df.columns):
            t.rows[0].cells[j].text = str(col)
        for i, (_, row) in enumerate(df.iterrows()):
            for j, val in enumerate(row):
                t.rows[i+1].cells[j].text = f'{val:.4f}' if isinstance(val, float) else str(val)
        sd.add_paragraph()

sd.add_heading('eTable 2. Metal Mixture Models', level=2)
mixture_path = TBL / 'logistic_core_model.csv'
if mixture_path.exists():
    df = pd.read_csv(str(mixture_path))
    sd.add_paragraph('Core model (without race):', style='Normal')
    t = sd.add_table(rows=len(df)+1, cols=len(df.columns))
    t.style = 'Light Grid Accent 1'
    for j, col in enumerate(df.columns):
        t.rows[0].cells[j].text = str(col)
    for i, (_, row) in enumerate(df.iterrows()):
        for j, val in enumerate(row):
            t.rows[i+1].cells[j].text = f'{val:.4f}' if isinstance(val, float) else str(val)

sd.add_paragraph()
sd.add_heading('eTable 3. Feature Importance and SHAP Comparison', level=2)
fi_path = TBL / 'feature_importance_enhanced.csv'
if fi_path.exists():
    df = pd.read_csv(str(fi_path))
    t = sd.add_table(rows=len(df)+1, cols=len(df.columns))
    t.style = 'Light Grid Accent 1'
    for j, col in enumerate(df.columns):
        t.rows[0].cells[j].text = str(col)
    for i, (_, row) in enumerate(df.iterrows()):
        for j, val in enumerate(row):
            t.rows[i+1].cells[j].text = f'{val:.4f}' if isinstance(val, float) else str(val)

supp_path = OUT / 'Supplementary_Materials.docx'
sd.save(str(supp_path))
print(f'Supplementary: {supp_path}')

# ── STROBE CHECKLIST ──
st = Document()
st_style = st.styles['Normal']
st_style.font.name = 'Times New Roman'
st_style.font.size = Pt(12)

st.add_heading('STROBE Statement — Checklist for Cross-Sectional Studies', level=1)
st.add_paragraph('Manuscript: Blood Lead, Mercury, and Cadmium Levels in Relation to Gestational Diabetes Mellitus: A Cross-Sectional Study and Machine Learning Analysis of NHANES 1999-2018')
st.add_paragraph('')

strobe_items = [
    ('Title/Abstract', '1', '(a) Indicate the study design in the title/abstract', 'Title and Abstract (p1-2)', ''),
    ('Title/Abstract', '2', 'Provide a structured abstract with Background, Methods, Results, Conclusions', 'Abstract (p1)', ''),
    ('Background', '3', 'Explain scientific background and rationale', 'Background (p3-4)', ''),
    ('Background', '4', 'State specific objectives/hypotheses', 'Background (end, p4)', ''),
    ('Methods', '5', 'Present key study design elements', 'Methods - Study population', ''),
    ('Methods', '6', 'Describe setting, locations, dates', 'Methods - Study population', ''),
    ('Methods', '7', 'Give eligibility criteria, sources/selection of participants', 'Methods - Study population', ''),
    ('Methods', '8', 'Clearly define all outcomes, exposures, predictors', 'Methods - GDM, Metals, Covariates', ''),
    ('Methods', '9', 'Describe data sources/measurement methods', 'Methods - Blood metal measurements', ''),
    ('Methods', '10', 'Describe efforts to address bias', 'Methods - Statistical analysis (survey weights, robust SE)', ''),
    ('Methods', '11', 'Explain how sample size was determined', 'NHANES pre-determined sample (n=10,979 after criteria)', ''),
    ('Methods', '12', 'Explain quantitative variables and how handled', 'Methods - Covariates, Statistical analysis (log-transformed metals)', ''),
    ('Methods', '13', '(a) Describe all statistical methods; (b) Subgroup analyses; (c) Missing data; (d) Survey weights', 'Methods - Statistical analysis, Race-stratified analysis', ''),
    ('Results', '14', 'Report numbers of individuals at each stage', 'Results - Study population (n=10,979)', ''),
    ('Results', '15', 'Report baseline characteristics (Table 1)', 'Table 1', ''),
    ('Results', '16', 'Report numbers of outcome events/summary measures', 'Results - Study population (3,527 GDM cases)', ''),
    ('Results', '17', 'Give unadjusted and adjusted estimates (OR with CI)', 'Table 2, Results - Logistic regression', ''),
    ('Results', '18', 'Report subgroup analyses', 'Results - Race-stratified analysis', ''),
    ('Results', '19', 'Summarize key results with reference to study objectives', 'Discussion - first paragraph', ''),
    ('Discussion', '20', 'Discuss limitations, potential bias, generalizability', 'Discussion - Limitations', ''),
    ('Discussion', '21', 'Give cautious interpretation considering limitations', 'Discussion (throughout)', ''),
    ('Discussion', '22', 'Discuss external validity (generalizability)', 'Discussion - Limitations', ''),
    ('Other', '23', 'Give source of funding', 'Declarations - Funding', ''),
]

t = st.add_table(rows=len(strobe_items)+1, cols=5)
t.style = 'Light Grid Accent 1'
headers = ['Section', 'Item', 'Recommendation', 'Manuscript Location', 'Page']
for j, h in enumerate(headers):
    t.rows[0].cells[j].text = h
for i, (section, item, rec, loc, pg) in enumerate(strobe_items):
    t.rows[i+1].cells[0].text = section
    t.rows[i+1].cells[1].text = item
    t.rows[i+1].cells[2].text = rec
    t.rows[i+1].cells[3].text = loc
    t.rows[i+1].cells[4].text = pg

strobe_path = OUT / 'STROBE_Checklist.docx'
st.save(str(strobe_path))
print(f'STROBE: {strobe_path}')

print('\n=== ALL GENERATED ===')
print(f'  1. {cover_path}')
print(f'  2. {supp_path}')
print(f'  3. {strobe_path}')
